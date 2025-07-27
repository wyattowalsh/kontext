"""
Document Processing Module for Kontext.

Handles ingestion and processing of various document formats using the docling library
for LLM-ready context extraction with async processing and progress tracking.
"""

import asyncio
import io
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
import polars as pl
from datetime import datetime
import streamlit as st

# Document processing libraries
from docling import DocumentConverter
from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import PdfPipelineOptions
from docling.document_converter import PdfFormatOption
import PyPDF2
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import markdown

# Async and progress
import anyio
from tqdm.asyncio import tqdm
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, TaskProgressColumn

# Logging
from loguru import logger

# Local imports
from config.settings import AppConfig
from utils.validators import validate_file_type, validate_file_size

console = Console()

class DocumentProcessor:
    """
    Advanced document processor with async capabilities and LLM-ready output.
    
    Supports multiple document formats and provides structured extraction
    with configurable chunking and metadata extraction.
    """
    
    def __init__(self, config: AppConfig):
        """
        Initialize the document processor.
        
        Args:
            config: Application configuration
        """
        self.config = config
        self.doc_config = config.document_processing
        
        # Initialize docling converter with optimized settings
        self.converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(
                    pipeline_options=PdfPipelineOptions(
                        do_ocr=True,
                        do_table_structure=self.doc_config.extract_tables,
                        table_structure_options={
                            "do_cell_matching": True,
                        }
                    )
                )
            }
        )
        
        logger.info("DocumentProcessor initialized with docling converter")
    
    async def process_files_async(
        self, 
        files: List[Any], 
        options: Dict[str, Any],
        progress_placeholder: Optional[Any] = None
    ) -> List[Dict[str, Any]]:
        """
        Process multiple files asynchronously with progress tracking.
        
        Args:
            files: List of uploaded files from Streamlit
            options: Processing options (chunk_size, extract_tables, etc.)
            progress_placeholder: Streamlit placeholder for progress updates
            
        Returns:
            List of processing results with extracted content and metadata
        """
        logger.info(f"Starting async processing of {len(files)} files")
        
        results = []
        
        # Create progress tracking
        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            BarColumn(),
            TaskProgressColumn(),
            console=console
        ) as progress:
            
            task = progress.add_task("Processing documents...", total=len(files))
            
            # Process files with controlled concurrency
            semaphore = asyncio.Semaphore(self.doc_config.max_workers)
            
            async def process_single_file(file_obj) -> Dict[str, Any]:
                async with semaphore:
                    try:
                        result = await self._process_single_file_async(file_obj, options)
                        progress.advance(task)
                        
                        # Update Streamlit progress if available
                        if progress_placeholder:
                            current_progress = len(results) / len(files)
                            progress_placeholder.progress(current_progress)
                            
                        return result
                    except Exception as e:
                        logger.error(f"Error processing file {file_obj.name}: {e}")
                        progress.advance(task)
                        return {
                            'filename': file_obj.name,
                            'status': 'error',
                            'error': str(e),
                            'processed_at': datetime.now().isoformat()
                        }
            
            # Execute all file processing tasks
            tasks = [process_single_file(file_obj) for file_obj in files]
            results = await asyncio.gather(*tasks, return_exceptions=True)
            
            # Filter out exceptions and convert to proper results
            processed_results = []
            for result in results:
                if isinstance(result, Exception):
                    logger.error(f"Task exception: {result}")
                    processed_results.append({
                        'status': 'error',
                        'error': str(result),
                        'processed_at': datetime.now().isoformat()
                    })
                else:
                    processed_results.append(result)
        
        logger.info(f"Completed processing {len(processed_results)} files")
        return processed_results
    
    async def _process_single_file_async(
        self, 
        file_obj: Any, 
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Process a single file asynchronously.
        
        Args:
            file_obj: Streamlit uploaded file object
            options: Processing options
            
        Returns:
            Processing result with extracted content and metadata
        """
        filename = file_obj.name
        file_extension = Path(filename).suffix.lower().lstrip('.')
        
        logger.debug(f"Processing file: {filename} (type: {file_extension})")
        
        # Validate file
        if not validate_file_type(filename, self.doc_config.supported_formats):
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        if not validate_file_size(file_obj, self.doc_config.max_file_size_mb):
            raise ValueError(f"File too large (max: {self.doc_config.max_file_size_mb}MB)")
        
        # Read file content
        file_content = file_obj.read()
        file_obj.seek(0)  # Reset file pointer
        
        # Process based on file type
        if file_extension == 'pdf':
            result = await self._process_pdf_async(file_content, filename, options)
        elif file_extension == 'docx':
            result = await self._process_docx_async(file_content, filename, options)
        elif file_extension in ['txt', 'md']:
            result = await self._process_text_async(file_content, filename, options)
        elif file_extension == 'html':
            result = await self._process_html_async(file_content, filename, options)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        # Add common metadata
        result.update({
            'filename': filename,
            'file_type': file_extension,
            'file_size_bytes': len(file_content),
            'processed_at': datetime.now().isoformat(),
            'processing_options': options,
            'status': 'success'
        })
        
        return result
    
    async def _process_pdf_async(
        self, 
        content: bytes, 
        filename: str, 
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Process PDF file using docling for advanced extraction."""
        logger.debug(f"Processing PDF: {filename}")
        
        try:
            # Use docling for advanced PDF processing
            # Create temporary file-like object
            pdf_stream = io.BytesIO(content)
            
            # Convert using docling
            result = self.converter.convert(pdf_stream)
            
            # Extract structured content
            extracted_text = result.document.export_to_markdown()
            
            # Extract tables if requested
            tables = []
            if options.get('extract_tables', True) and hasattr(result.document, 'tables'):
                for table in result.document.tables:
                    tables.append({
                        'table_data': table.export_to_dataframe().to_dict() if hasattr(table, 'export_to_dataframe') else str(table),
                        'caption': getattr(table, 'caption', ''),
                        'page': getattr(table, 'page', 0)
                    })
            
            # Extract images if requested
            images = []
            if options.get('extract_images', False) and hasattr(result.document, 'pictures'):
                for img in result.document.pictures:
                    images.append({
                        'description': getattr(img, 'description', ''),
                        'page': getattr(img, 'page', 0),
                        'bbox': getattr(img, 'bbox', None)
                    })
            
            # Chunk text if requested
            chunks = self._chunk_text(extracted_text, options.get('chunk_size', 512))
            
            return {
                'extracted_text': extracted_text,
                'chunks': chunks,
                'tables': tables,
                'images': images,
                'page_count': getattr(result.document, 'page_count', 0),
                'word_count': len(extracted_text.split()),
                'char_count': len(extracted_text)
            }
            
        except Exception as e:
            logger.error(f"Docling processing failed for {filename}: {e}")
            # Fallback to PyPDF2
            return await self._process_pdf_fallback_async(content, filename, options)
    
    async def _process_pdf_fallback_async(
        self, 
        content: bytes, 
        filename: str, 
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Fallback PDF processing using PyPDF2."""
        logger.debug(f"Using PyPDF2 fallback for: {filename}")
        
        pdf_stream = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_stream)
        
        extracted_text = ""
        for page in pdf_reader.pages:
            extracted_text += page.extract_text() + "\n"
        
        chunks = self._chunk_text(extracted_text, options.get('chunk_size', 512))
        
        return {
            'extracted_text': extracted_text,
            'chunks': chunks,
            'tables': [],
            'images': [],
            'page_count': len(pdf_reader.pages),
            'word_count': len(extracted_text.split()),
            'char_count': len(extracted_text)
        }
    
    async def _process_docx_async(
        self, 
        content: bytes, 
        filename: str, 
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Process DOCX file."""
        logger.debug(f"Processing DOCX: {filename}")
        
        docx_stream = io.BytesIO(content)
        doc = DocxDocument(docx_stream)
        
        # Extract text from paragraphs
        extracted_text = ""
        for paragraph in doc.paragraphs:
            extracted_text += paragraph.text + "\n"
        
        # Extract tables if requested
        tables = []
        if options.get('extract_tables', True):
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append({
                    'table_data': table_data,
                    'caption': '',
                    'rows': len(table_data),
                    'cols': len(table_data[0]) if table_data else 0
                })
        
        chunks = self._chunk_text(extracted_text, options.get('chunk_size', 512))
        
        return {
            'extracted_text': extracted_text,
            'chunks': chunks,
            'tables': tables,
            'images': [],
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
            'word_count': len(extracted_text.split()),
            'char_count': len(extracted_text)
        }
    
    async def _process_text_async(
        self, 
        content: bytes, 
        filename: str, 
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Process plain text or Markdown file."""
        logger.debug(f"Processing text file: {filename}")
        
        # Decode text content
        try:
            text_content = content.decode('utf-8')
        except UnicodeDecodeError:
            text_content = content.decode('latin-1')
        
        # Convert Markdown to HTML if it's a .md file
        html_content = ""
        if filename.lower().endswith('.md'):
            html_content = markdown.markdown(text_content, extensions=['tables', 'fenced_code'])
        
        chunks = self._chunk_text(text_content, options.get('chunk_size', 512))
        
        return {
            'extracted_text': text_content,
            'html_content': html_content,
            'chunks': chunks,
            'tables': [],
            'images': [],
            'line_count': len(text_content.split('\n')),
            'word_count': len(text_content.split()),
            'char_count': len(text_content)
        }
    
    async def _process_html_async(
        self, 
        content: bytes, 
        filename: str, 
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Process HTML file."""
        logger.debug(f"Processing HTML: {filename}")
        
        # Decode HTML content
        try:
            html_content = content.decode('utf-8')
        except UnicodeDecodeError:
            html_content = content.decode('latin-1')
        
        # Parse with BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Extract text content
        extracted_text = soup.get_text(separator='\n', strip=True)
        
        # Extract tables if requested
        tables = []
        if options.get('extract_tables', True):
            for table in soup.find_all('table'):
                table_data = []
                for row in table.find_all('tr'):
                    row_data = [cell.get_text(strip=True) for cell in row.find_all(['td', 'th'])]
                    if row_data:
                        table_data.append(row_data)
                
                if table_data:
                    tables.append({
                        'table_data': table_data,
                        'caption': table.find('caption').get_text(strip=True) if table.find('caption') else '',
                        'rows': len(table_data),
                        'cols': len(table_data[0]) if table_data else 0
                    })
        
        # Extract images if requested
        images = []
        if options.get('extract_images', False):
            for img in soup.find_all('img'):
                images.append({
                    'src': img.get('src', ''),
                    'alt': img.get('alt', ''),
                    'title': img.get('title', '')
                })
        
        chunks = self._chunk_text(extracted_text, options.get('chunk_size', 512))
        
        return {
            'extracted_text': extracted_text,
            'html_content': html_content,
            'chunks': chunks,
            'tables': tables,
            'images': images,
            'title': soup.find('title').get_text(strip=True) if soup.find('title') else '',
            'word_count': len(extracted_text.split()),
            'char_count': len(extracted_text)
        }
    
    def _chunk_text(self, text: str, chunk_size: int) -> List[Dict[str, Any]]:
        """
        Split text into chunks for LLM processing.
        
        Args:
            text: Text to chunk
            chunk_size: Target chunk size in tokens (approximated by words)
            
        Returns:
            List of text chunks with metadata
        """
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size):
            chunk_words = words[i:i + chunk_size]
            chunk_text = ' '.join(chunk_words)
            
            chunks.append({
                'chunk_id': len(chunks),
                'text': chunk_text,
                'word_count': len(chunk_words),
                'char_count': len(chunk_text),
                'start_word': i,
                'end_word': min(i + chunk_size, len(words))
            })
        
        return chunks
    
    def export_results(
        self, 
        results: List[Dict[str, Any]], 
        format: str = 'json',
        output_path: Optional[str] = None
    ) -> str:
        """
        Export processing results to specified format.
        
        Args:
            results: Processing results to export
            format: Export format ('json', 'csv', 'parquet')
            output_path: Output file path
            
        Returns:
            Path to exported file
        """
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"kontext_processing_results_{timestamp}.{format}"
        
        # Convert to Polars DataFrame for consistent export
        df = pl.DataFrame(results)
        
        if format == 'json':
            df.write_json(output_path)
        elif format == 'csv':
            df.write_csv(output_path)
        elif format == 'parquet':
            df.write_parquet(output_path)
        else:
            raise ValueError(f"Unsupported export format: {format}")
        
        logger.info(f"Results exported to {output_path}")
        return output_path